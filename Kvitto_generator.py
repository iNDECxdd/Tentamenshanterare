import re
import sys
import os
from docx import Document
import pyperclip


def extrahera_data(pdf_text: str):
    data = {
        'Personnummer': '',
        'Förnamn': '',
        'Efternamn': '',
        'Adress': '',
        'Postnummer': '',
        'Ort': '',
        'Kurskod': '',
        'Lärosäte': '',
        'Datum': '',
    }

    lines = (pdf_text or "").split('\n')
    current_key = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if re.match(r'^Personnummer\s+', line, re.IGNORECASE):
            current_key = 'Personnummer'
            value = re.split(r'\s+', line, maxsplit=1)[-1].strip()
            data[current_key] = value

        elif re.match(r'^Förnamn\s+', line, re.IGNORECASE):
            current_key = 'Förnamn'
            value = re.split(r'\s+', line, maxsplit=1)[-1].strip()
            data[current_key] = value

        elif re.match(r'^Efternamn\s+', line, re.IGNORECASE):
            current_key = 'Efternamn'
            value = re.split(r'\s+', line, maxsplit=1)[-1].strip()
            data[current_key] = value

        elif re.match(r'^Adress\s+', line, re.IGNORECASE):
            current_key = 'Adress'
            value = re.split(r'\s+', line, maxsplit=1)[-1].strip()
            data[current_key] = value

        elif re.match(r'^Postnummer\s+', line, re.IGNORECASE):
            current_key = 'Postnummer'
            value = re.split(r'\s+', line, maxsplit=1)[-1].strip()
            data[current_key] = value

        elif re.match(r'^Ort\s+', line, re.IGNORECASE):
            current_key = 'Ort'
            value = re.split(r'\s+', line, maxsplit=1)[-1].strip()
            data[current_key] = value

        elif re.match(r'^Kurskod\s+', line, re.IGNORECASE):
            current_key = 'Kurskod'
            value = re.split(r'\s+', line, maxsplit=1)[-1].strip()
            data[current_key] = value

        elif re.match(r'^Högskola, universitet eller annat lärosäte\s+', line, re.IGNORECASE):
            current_key = 'Lärosäte'
            m = re.match(r'^Högskola, universitet eller annat lärosäte\s+(.*)$', line, re.IGNORECASE)
            value = (m.group(1) if m else "").strip()

            lower_value = re.sub(r'\s+', ' ', value.lower())

            if re.search(r'\bkarlstads?\s+universitet\b|\bkau\b|\bkarlstad\b', lower_value):
                data[current_key] = "KAU"
            elif re.search(r'\bmittuniversitetet\b|\bmiun\b', lower_value):
                data[current_key] = "MIUN"
            else:
                data[current_key] = value

        elif "datum för tentamen" in line.lower():
            match = re.search(r'datum för tentamen\s+(.+)', line, re.IGNORECASE)
            if match:
                data['Datum'] = match.group(1).strip()

        # Om adress fortsätter på nästa rad (tills en annan nyckel kommer)
        elif current_key == 'Adress' and not any(
            re.match(fr'^{key}\s+', line, re.IGNORECASE)
            for key in data if key != 'Adress'
        ):
            data['Adress'] += '\n' + line

    data['Adress'] = data['Adress'].split('Postnummer')[0].strip()

    return [
        data['Personnummer'],
        data['Förnamn'],
        data['Efternamn'],
        data['Adress'],
        data['Postnummer'],
        data['Ort'],
        data['Kurskod'],
        data['Lärosäte'],
        data['Datum'],
    ]

def print_word_docx(file_path: str, copies: int = 2):
    """
    Skriver ut en .docx via Word (Windows). Kräver pywin32:
    pip install pywin32
    """
    import os
    try:
        import win32com.client  # type: ignore
    except Exception as e:
        print(f"[VARNING] Kan inte skriva ut (saknar pywin32). Installera med: pip install pywin32\n{e}")
        return

    abs_path = os.path.abspath(file_path)

    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(abs_path, ReadOnly=True)

        # ✅ Robust: skriv ut 1 kopia flera gånger (vissa drivers ignorerar Copies>1)
        for _ in range(int(copies)):
            doc.PrintOut(Copies=1, Background=False)

    except Exception as e:
        print(f"[FEL] Utskrift misslyckades: {e}")
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass


def ersätt_platshållare(dokument: Document, delar):
    # Ersätt i paragrafer
    for paragraph in dokument.paragraphs:
        text = paragraph.text
        for i in range(9):
            ph = f'{{{i+1}}}'
            if ph in text:
                text = text.replace(ph, delar[i])
        if text != paragraph.text:
            paragraph.clear()
            paragraph.add_run(text)

    # Ersätt i tabeller
    for table in dokument.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    for i in range(9):
                        ph = f'{{{i+1}}}'
                        if ph in text:
                            text = text.replace(ph, delar[i])
                    if text != paragraph.text:
                        paragraph.clear()
                        paragraph.add_run(text)


def _read_input_text() -> str:
    """
    Om en fil skickas som argument: läs den (UTF-8).
    Annars: använd urklipp (bakåtkompatibelt med din gamla workflow).
    """
    if len(sys.argv) >= 2 and sys.argv[1]:
        path = sys.argv[1]
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    return pyperclip.paste()


def huvudfunktion():
    pdf_text = _read_input_text()
    delar = extrahera_data(pdf_text)

    doc = Document('Tentamenskvitto.docx')
    ersätt_platshållare(doc, delar)
    doc.save('resultat.docx')
    print("Klart! Formateringen är bevarad i 'resultat.docx'.")

    print_word_docx("resultat.docx", copies=2)


if __name__ == "__main__":
    huvudfunktion()
